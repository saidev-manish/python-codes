from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create a new Word document
doc = Document()

# Helper functions
def add_heading_center(text, size=14, bold=True):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_paragraph_justified(text, size=12):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in p.runs:
        run.font.size = Pt(size)

# Title Page
add_heading_center("HYDERABAD INSTITUTE OF TECHNOLOGY AND MANAGEMENT", 14, True)
add_heading_center("Department of Computer Science and Engineering (Data Science)", 12, True)
add_heading_center("(UGC Autonomous, Accredited by NAAC (A+), NBA)", 11, False)
doc.add_paragraph()
add_heading_center("A PBL PROJECT REPORT ON", 13, True)
add_heading_center("“ULTIMATE UTILITY & SCIENTIFIC CALCULATOR”", 14, True)
add_paragraph_justified("In partial fulfilment of the requirements for the award of the degree of Bachelor of Technology in COMPUTER SCIENCE AND ENGINEERING (DATA SCIENCE)")
doc.add_paragraph()
add_paragraph_justified("Submitted by")
add_paragraph_justified("Kavali Ankitha (24E51A6776)\nMydhari Ishwar Saidevmanish (24E51A67C3)\nMandepudi Lahari (24E51A67A6)\nJuvvanapudi Bhavishya (24E51A6764)")
doc.add_paragraph()
add_paragraph_justified("Under the Esteemed Guidance of\nMr. K. Ravi Kumar (Associate Professor)")
add_paragraph_justified("Department of Computer Science and Engineering (Data Science)\nHyderabad Institute of Technology and Management\nGowdavelly (Village), Medchal, Hyderabad, Telangana – 501401\n2025")

# Certificate
doc.add_page_break()
add_heading_center("CERTIFICATE", 13, True)
add_paragraph_justified(
    "This is to certify that the Project work entitled “ULTIMATE UTILITY & SCIENTIFIC CALCULATOR” "
    "is a bonafide work carried out by Kavali Ankitha (24E51A6776), Mydhari Ishwar Saidevmanish (24E51A67C3), "
    "Mandepudi Lahari (24E51A67A6), and Juvvanapudi Bhavishya (24E51A6764) in partial fulfilment of the requirements "
    "for the award of the degree of Bachelor of Technology in Computer Science and Engineering (Data Science) by the "
    "Jawaharlal Nehru Technological University, Hyderabad, during the academic year 2025–2026."
)
add_paragraph_justified(
    "The matter embodied in this project report has not been submitted to any other university or institute for the award of any degree or diploma."
)
add_paragraph_justified(
    "Internal Supervisor: Mr. K. Ravi Kumar (Associate Professor)\n"
    "Program Head: Mr. Bhaskar Das (Associate Professor)\n"
    "Head of Department: Dr. Kolluri David Raju (Professor & HoD)\n"
    "External Examiner: ____________________"
)

# Declaration
doc.add_page_break()
add_heading_center("DECLARATION", 13, True)
add_paragraph_justified(
    "We, Kavali Ankitha (24E51A6776), Mydhari Ishwar Saidevmanish (24E51A67C3), Mandepudi Lahari (24E51A67A6), "
    "and Juvvanapudi Bhavishya (24E51A6764), students of Bachelor of Technology in Computer Science and Engineering (Data Science), "
    "session 2024–2025, Hyderabad Institute of Technology and Management, Gowdavelly, Hyderabad, Telangana State, hereby declare that "
    "the work presented in this project entitled “ULTIMATE UTILITY & SCIENTIFIC CALCULATOR” is the outcome of our own bonafide effort and is correct to the best of our knowledge.\n\n"
    "This project work has been carried out adhering to engineering ethics. It contains no material previously published or written by any other person, "
    "nor material accepted for the award of any other degree or diploma of the university or any other institute of higher learning, except where due acknowledgment has been made in the text."
)
add_paragraph_justified(
    "KAVALI ANKITHA\t\t24E51A6776\n"
    "MYDHARI ISHWAR SAIDEVMANISH\t24E51A67C3\n"
    "MANDEPUDI LAHARI\t\t24E51A67A6\n"
    "JUVVANAPUDI BHAVISHYA\t24E51A6764"
)

# Acknowledgment
doc.add_page_break()
add_heading_center("ACKNOWLEDGMENT", 13, True)
add_paragraph_justified(
    "An endeavour of a long period can only be successful with the guidance, motivation, and encouragement of many well-wishers.\n\n"
    "We express our deep sense of gratitude to our Chairman, Sri Arutla Prashanth, for providing us with all the facilities and resources required to successfully complete our project work.\n\n"
    "We extend our sincere thanks to our Honourable Principal, Dr. Arvind, whose inspiring words and encouragement have greatly motivated us and for providing us with this opportunity to undertake our major project.\n\n"
    "We are highly grateful to our Head of the Department, Dr. Kolluri David Raju, our Program Head, Mr. Bhaskar Das, and our B.Tech Project Coordinator, Mr. K. Ravi Kumar (Associate Professor), for their valuable guidance, continuous support, and motivation throughout the course of this project.\n\n"
    "We would like to make a special mention of our Internal Supervisor, Mr. K. Ravi Kumar, for his invaluable technical guidance, constant encouragement, and immense support, which helped us in the successful completion of our project “Ultimate Utility & Scientific Calculator”.\n\n"
    "We also extend our heartfelt thanks to all the Departmental Committee (D.C.) and Project Review Committee (P.R.C.) members, as well as the non-teaching staff, for their cooperation and assistance during our project work.\n\n"
    "Finally, we express our deep gratitude to our family members and friends for their continuous encouragement, moral support, and care, which have been our greatest strength throughout this journey."
)

# Abstract
doc.add_page_break()
add_heading_center("ABSTRACT", 13, True)
add_paragraph_justified(
    "The “Ultimate Utility & Scientific Calculator” is a Python-based console application that performs both basic arithmetic and advanced scientific operations, along with academic utilities like SGPA and CGPA calculation. "
    "This project demonstrates key programming concepts including modular design using functions, user input handling and validation, mathematical computations using Python’s math library, and logical implementation of number theory and temperature conversion. "
    "This versatile calculator provides an easy and efficient solution for students and professionals to perform diverse mathematical and academic operations through a single unified system."
)

# Conclusion & References
doc.add_page_break()
add_heading_center("CONCLUSION", 13, True)
add_paragraph_justified(
    "The “Ultimate Utility & Scientific Calculator” project demonstrates a practical understanding of Python programming concepts including modularization, control structures, and mathematical operations. "
    "This project serves as a foundation for developing more advanced applications and can be extended into a Graphical User Interface (GUI) or web-based calculator in future work."
)

add_heading_center("REFERENCES", 13, True)
add_paragraph_justified(
    "1. Python Documentation: https://docs.python.org/3/\n"
    "2. Math Library Reference: https://docs.python.org/3/library/math.html\n"
    "3. Stack Overflow & GeeksforGeeks Articles on Python Math Functions"
)

# Save the document
import os
downloads_path = os.path.join(os.path.expanduser('~'), 'Downloads')
file_path = os.path.join(downloads_path, "Ultimate_Utility_Scientific_Calculator_Project_Report.docx")
doc.save(file_path)

print(f"Document saved successfully as: {file_path}")
